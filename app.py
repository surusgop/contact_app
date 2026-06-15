import os
import io
import json
import base64
import threading
import requests
import urllib3
import pandas as pd
from datetime import datetime, timezone
from flask import Flask, render_template, request, jsonify, redirect, url_for, session
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from authlib.integrations.flask_client import OAuth
from dotenv import load_dotenv
from databricks.sdk import WorkspaceClient
from databricks.sdk.service.sql import StatementParameterListItem, StatementState

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
_original_send = requests.Session.send
def _send_no_verify(self, *args, **kwargs):
    kwargs['verify'] = False
    return _original_send(self, *args, **kwargs)
requests.Session.send = _send_no_verify

load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "dev-secret-change-me")

login_manager = LoginManager(app)
login_manager.login_view = "login"

oauth = OAuth(app)
google = oauth.register(
    name="google",
    client_id=os.getenv("GOOGLE_CLIENT_ID"),
    client_secret=os.getenv("GOOGLE_CLIENT_SECRET"),
    server_metadata_url="https://accounts.google.com/.well-known/openid-configuration",
    client_kwargs={"scope": "openid email profile"},
)

class User(UserMixin):
    def __init__(self, id, email, name, picture):
        self.id = id
        self.email = email
        self.name = name
        self.picture = picture

_users = {}

@login_manager.user_loader
def load_user(user_id):
    return _users.get(user_id)

_db = None
def get_db():
    global _db
    if _db is None:
        _db = WorkspaceClient(
            host=os.getenv("DATABRICKS_HOST"),
            token=os.getenv("DATABRICKS_TOKEN"),
        )
    return _db

WAREHOUSE_ID = os.getenv("DATABRICKS_WAREHOUSE_ID", "")

CONTACT_METHODS = [
    "delivery", "door_knock", "email", "email_blast", "face_to_face",
    "facebook", "meeting", "phone_call", "robocall", "snail_mail",
    "text", "text_1to1", "text_blast", "tweet", "video_call",
    "webinar", "linkedin", "other",
]
CONTACT_STATUSES = [
    "answered", "bad_info", "left_message", "meaningful_interaction",
    "send_information", "not_interested", "no_answer", "refused",
    "inaccessible", "other",
]

_NICKNAME_GROUPS = [
    # Male
    ("harry", "harrison", "harold", "henry", "hank", "hal", "harris", "harvey"),
    ("bill", "billy", "william", "will", "willy", "liam", "willem"),
    ("bob", "bobby", "robert", "rob", "robbie", "bert", "bertie"),
    ("jim", "jimmy", "james", "jamie", "jake"),
    ("joe", "joey", "joseph", "jose"),
    ("jack", "john", "johnny", "jon", "jonathan", "jacky"),
    ("mike", "mikey", "michael", "mick", "mickey"),
    ("dave", "david", "davy", "davey"),
    ("tom", "tommy", "thomas"),
    ("rick", "ricky", "richard", "dick", "rich", "richie"),
    ("nick", "nicky", "nicholas", "nicolas", "nico"),
    ("chris", "christopher", "christian", "cris", "kris"),
    ("dan", "danny", "daniel", "dani"),
    ("al", "albert", "alfred", "alan", "ali", "alfie"),
    ("ed", "eddie", "edward", "edgar", "ned", "ted", "ward"),
    ("frank", "frankie", "francis", "fran"),
    ("fred", "freddy", "frederick", "fritz"),
    ("jerry", "gerald", "jerome", "gerry"),
    ("ken", "kenny", "kenneth"),
    ("larry", "lawrence", "lars", "laurence"),
    ("pete", "peter", "petey"),
    ("ray", "raymond"),
    ("ron", "ronnie", "ronald"),
    ("steve", "steven", "stephen", "stevo"),
    ("tim", "timmy", "timothy"),
    ("tony", "anthony", "ant"),
    ("andy", "andrew", "drew"),
    ("ben", "benny", "benjamin"),
    ("brad", "bradley", "bradford"),
    ("charlie", "charles", "chuck", "chas", "carl", "carlos"),
    ("don", "donald", "donnie"),
    ("doug", "douglas", "dougie"),
    ("greg", "gregory", "gregg"),
    ("jeff", "jeffrey", "geoff", "geoffrey"),
    ("josh", "joshua"),
    ("matt", "matthew", "matty"),
    ("max", "maxwell", "maximilian", "maxine", "maxime"),
    ("nat", "nate", "nathan", "nathaniel"),
    ("neil", "neal", "nigel"),
    ("norm", "norman"),
    ("ollie", "oliver"),
    ("phil", "phillip", "philip"),
    ("rod", "rodney", "roderick"),
    ("russ", "russell"),
    ("sam", "samuel", "sammy"),
    ("stan", "stanley", "stanford"),
    ("stu", "stuart", "stewart"),
    ("theo", "theodore", "teddy"),
    ("vince", "vincent", "vinny"),
    ("walt", "walter", "wally"),
    ("wes", "wesley", "weston"),
    ("woody", "woodrow"),
    ("zach", "zachary", "zack", "zak"),
    ("alex", "alexander", "alec", "lex"),
    ("art", "arthur", "artie"),
    ("bart", "barton", "bartholomew"),
    ("bud", "buddy"),
    ("cal", "calvin", "caleb"),
    ("chet", "chester"),
    ("clint", "clinton"),
    ("cy", "cyrus"),
    ("denny", "dennis", "denis"),
    ("dom", "dominic", "dominick"),
    ("dusty", "dustin"),
    ("eli", "elijah", "elias"),
    ("ernie", "ernest"),
    ("gabe", "gabriel"),
    ("gene", "eugene"),
    ("gil", "gilbert"),
    ("gus", "augustus", "angus", "gustav"),
    ("herb", "herbert"),
    ("ike", "isaac"),
    ("jay", "james", "jason", "jacob"),
    ("jeff", "jeffrey"),
    ("leo", "leonard", "leon"),
    ("les", "leslie", "lester"),
    ("lew", "lewis", "louis", "lou"),
    ("mac", "malcolm", "mack"),
    ("manny", "manuel", "emmanuel"),
    ("marty", "martin"),
    ("mel", "melvin"),
    ("mitch", "mitchell"),
    ("monty", "montgomery"),
    ("mort", "morton", "mortimer"),
    ("moe", "morris", "moses"),
    ("ozzy", "oscar", "oswald"),
    ("percy", "percival"),
    ("reg", "reginald"),
    ("rex", "reginald"),
    ("rudy", "rudolph", "rudolf"),
    ("scott", "scotty"),
    ("sean", "shaun", "shawn"),
    ("sherm", "sherman"),
    ("sid", "sydney", "sylvester"),
    ("skip", "james", "john"),
    ("sol", "solomon"),
    ("sonny", "sylvester"),
    ("tad", "thaddeus"),
    ("trey", "trevor", "trenton"),
    ("zeke", "ezekiel"),
    ("rand", "randall", "randolph", "randy"),
    ("reggie", "reginald"),
    ("rocky", "rockwell"),
    ("ross", "roscoe"),
    ("dex", "dexter"),
    ("kurt", "curtis", "curt"),
    ("wade", "walden"),
    ("xander", "alexander"),
    ("emmett", "emmet"),
    ("glen", "glenn"),
    ("heath", "heathcliff"),
    ("jed", "jedidiah"),
    ("kip", "kipling"),
    ("lyle", "leland"),
    ("marsh", "marshall", "marcus"),
    ("murray", "muriel"),
    ("otis", "otto"),
    ("rad", "radley"),
    ("rush", "russell"),
    ("sly", "sylvester"),
    ("tanner", "tan"),
    # Female
    ("abby", "abigail", "abbey"),
    ("addie", "adelaide", "adeline", "ada"),
    ("aggie", "agnes"),
    ("allie", "alice", "allison", "alyssa"),
    ("angie", "angela", "angelina"),
    ("annie", "ann", "anna", "anne", "annette"),
    ("bea", "beatrice", "beatrix"),
    ("bella", "isabella", "arabella", "isabel"),
    ("bess", "bessie", "elizabeth"),
    ("beth", "elizabeth", "bethany"),
    ("betsy", "elizabeth"),
    ("betty", "elizabeth", "bette"),
    ("billie", "wilhelmina"),
    ("bonnie", "bonita"),
    ("bree", "brianna", "breanna", "sabrina"),
    ("brit", "brittany", "britney"),
    ("callie", "carolyn", "caroline", "calista"),
    ("carrie", "caroline", "carol", "carolyn", "carla", "carly"),
    ("cass", "cassandra", "cassidy"),
    ("cat", "catherine", "catalina", "caitlin"),
    ("cathy", "catherine", "kathleen", "kathryn"),
    ("cece", "cecelia", "cecilia"),
    ("cher", "cheryl", "sharon", "cherie"),
    ("cindy", "cynthia", "lucinda"),
    ("claire", "clara", "clarissa"),
    ("connie", "constance", "cornelia"),
    ("cora", "corinne", "cordelia"),
    ("debbie", "deborah", "deb", "debra"),
    ("dee", "deanna", "diana", "delia"),
    ("dora", "dorothy", "theodora", "dorothea"),
    ("dot", "dorothy"),
    ("ellie", "eleanor", "ellen", "elizabeth", "elena", "elaine"),
    ("emma", "emily", "emmeline"),
    ("eva", "evelyn", "evangeline", "eve"),
    ("flo", "florence", "flora"),
    ("frankie", "frances", "francesca", "francine"),
    ("gail", "abigail", "gayle"),
    ("ginny", "virginia", "geneva"),
    ("gwen", "gwendolyn", "gwyneth", "guinevere", "gwenna"),
    ("hattie", "harriet", "henrietta"),
    ("harriet", "harriet", "hattie"),
    ("jackie", "jacqueline"),
    ("janie", "jane", "janet", "jan"),
    ("jess", "jessica", "jessamine", "jessie"),
    ("jo", "josephine", "joan", "joanna"),
    ("josie", "josephine"),
    ("jules", "julia", "juliet", "julie"),
    ("kate", "katie", "kathy", "katherine", "catherine", "kathryn", "kay"),
    ("kim", "kimberly"),
    ("kitty", "katherine", "kathryn", "kathleen"),
    ("laurie", "laura", "lauren", "laurel", "lori"),
    ("lexi", "alexandra", "alexis"),
    ("libby", "elizabeth"),
    ("lilly", "lillian", "lily"),
    ("liz", "elizabeth", "lisa", "beth", "eliza"),
    ("lola", "dolores", "charlotte"),
    ("lori", "lorraine", "laura", "loretta"),
    ("lottie", "charlotte", "carlotta"),
    ("lu", "lucy", "lucia", "lucille", "louisa"),
    ("lucy", "lucia", "lucille", "lucinda"),
    ("mae", "mary", "margaret", "may"),
    ("maggie", "margaret", "magdalene", "mags"),
    ("mandy", "amanda", "miranda"),
    ("margie", "margaret", "margery", "marge"),
    ("meg", "margaret", "megan"),
    ("mel", "melissa", "melinda", "melanie"),
    ("millie", "mildred", "millicent", "emily", "camille", "amelia"),
    ("minnie", "wilhelmina", "minerva", "mina"),
    ("missy", "melissa"),
    ("molly", "mary", "margaret"),
    ("nan", "nancy", "ann", "nanette"),
    ("nell", "eleanor", "ellen", "helen", "nellie"),
    ("nora", "eleanor", "honora", "leonora"),
    ("pam", "pamela"),
    ("pat", "patricia", "patience"),
    ("patty", "patricia", "martha"),
    ("penny", "penelope"),
    ("peg", "peggy", "margaret"),
    ("polly", "mary", "paula", "pauline"),
    ("rita", "margaret", "margarita"),
    ("rose", "rosemary", "rosalyn", "rosalie"),
    ("rosie", "rosemary", "rose", "rosanna"),
    ("sadie", "sarah"),
    ("sally", "sarah"),
    ("sam", "samantha", "samara", "sammy"),
    ("sandy", "sandra", "cassandra"),
    ("sasha", "alexandra", "natasha"),
    ("sherry", "sharon", "charlotte"),
    ("sonya", "sophia", "sonia"),
    ("sophie", "sophia"),
    ("steph", "stephanie", "stefanie"),
    ("sue", "susan", "suzy", "susanna"),
    ("tammy", "tamara"),
    ("tess", "theresa", "teresa", "tessa"),
    ("tina", "christina", "valentina", "martina"),
    ("trudy", "gertrude"),
    ("tari", "terrance", "terry", "tarry", "carolyn"),
    ("val", "valerie", "valentine"),
    ("vicki", "victoria"),
    ("viv", "vivienne", "vivian"),
    ("winnie", "winifred", "winona"),
    ("becca", "rebecca", "becky"),
    ("cam", "cameron", "camille", "camilla"),
    ("chrissy", "christine", "christina"),
    ("dani", "danielle", "daniela"),
    ("elsa", "elizabeth", "eleanor"),
    ("gabby", "gabrielle", "gabriela"),
    ("hallie", "harriet"),
    ("jodi", "judith", "jody"),
    ("lena", "elena", "helen", "magdalena", "selena"),
    ("mia", "maria", "miriam", "amelia"),
    ("nicki", "nicole", "nicola"),
    ("nina", "annina", "antonina"),
    ("nola", "magnolia", "cornelia"),
    ("rae", "rachel", "renee"),
    ("roxy", "roxanne", "roxanna"),
    ("ruby", "rubina"),
    ("shelly", "michelle", "rochelle"),
    ("terri", "theresa", "teresa", "terrence"),
    ("tori", "victoria"),
    ("trish", "patricia", "tricia"),
    ("wendy", "gwendolyn"),
    ("barb", "barbara"),
    ("carol", "caroline", "carolyn"),
    ("linda", "belinda"),
    ("jen", "jenny", "jennifer"),
    # Gender-neutral
    ("alex", "alexander", "alexandra", "alexis"),
    ("ashley", "ash"),
    ("casey", "cassidy"),
    ("devon", "devonne"),
    ("drew", "andrew", "andrea"),
    ("hayden", "hayley"),
    ("jamie", "james", "jamison"),
    ("jordan", "jordy"),
    ("kelly", "kelsey"),
    ("lee", "leona", "leo"),
    ("leslie", "les"),
    ("morgan", "morgana"),
    ("quinn", "quincy"),
    ("riley", "reilly"),
    ("robin", "roberta", "robert"),
    ("rory", "aurora", "roderick"),
    ("ryan", "ryanne"),
    ("shannon", "shauna", "shana"),
    ("shelby", "sheila"),
    ("skyler", "sky"),
    ("stacy", "anastasia"),
    ("taylor", "tay"),
    ("tracey", "tracy", "theresa"),
    ("whitney", "whit"),
    ("terry", "theresa", "teresa", "terrence"),
    # Additional female names
    ("kit", "katherine", "kathryn", "kathleen", "kristin", "kristina"),
    ("gwen", "gwendolyn", "gwyneth", "guinevere", "gwenna"),
    ("teddy", "theodora", "theodore", "edwina", "edna"),
    ("winnie", "winifred", "winnifred", "winona"),
    ("gabby", "gabrielle", "gabriela", "gabriella", "gabi"),
    ("dottie", "dorothy", "dot"),
    ("elspeth", "elizabeth", "elspie"),
    ("flossie", "florence", "flora", "flo"),
    ("georgie", "georgina", "georgiana", "georgia"),
    ("hattie", "harriet", "henrietta", "hatty"),
    ("issy", "isabelle", "isabella", "isabel", "isobel"),
    ("jojo", "josephine", "joanna", "joanne"),
    ("lala", "laura", "lara", "larissa"),
    ("lettie", "letitia", "leticia", "violet"),
    ("lexie", "alexandra", "alexis", "lexi"),
    ("liddy", "lydia"),
    ("lilah", "delilah", "lillian", "lila"),
    ("lindy", "linda", "belinda", "melinda"),
    ("livvy", "olivia", "livia"),
    ("lottie", "charlotte", "carlotta"),
    ("lu", "louisa", "lucy", "lucille", "lucia", "lucinda", "louise", "lulu"),
    ("maddie", "madeleine", "madeline", "madelyn", "madison"),
    ("mae", "mary", "margaret", "may", "maeve"),
    ("maisie", "margaret", "mary", "miriam"),
    ("mally", "malinda", "melinda", "mallory"),
    ("marnie", "marina", "marianne", "maren"),
    ("mattie", "martha", "matilda"),
    ("maude", "madeleine", "magdalene"),
    ("midge", "margaret", "miriam"),
    ("milly", "mildred", "millicent", "camille", "amelia", "emily"),
    ("mindy", "melinda", "miranda", "amanda"),
    ("miri", "miriam", "mirielle"),
    ("nettie", "annette", "antoinette", "henrietta"),
    ("nicky", "nicole", "nicola", "nicolette"),
    ("nixie", "nicole", "nicola"),
    ("nonnie", "nora", "eleanor", "honora"),
    ("olga", "helga"),
    ("opal", "ophelia"),
    ("pippa", "philippa", "penelope"),
    ("posy", "josephine", "rose", "rosemary"),
    ("prissy", "priscilla"),
    ("queenie", "regina", "queen"),
    ("remi", "remy", "remington"),
    ("rena", "irene", "serena", "katerina", "renata"),
    ("retta", "henrietta", "loretta", "marietta"),
    ("ricki", "erica", "frederica", "rica"),
    ("rilla", "amarilla", "marilla"),
    ("romie", "rosemary", "rome"),
    ("roni", "veronica", "rhonda"),
    ("ronnie", "veronica", "rhonda", "rhona"),
    ("rosalie", "rose", "rosemary", "rosalyn"),
    ("roxie", "roxanne", "roxanna"),
    ("ruthie", "ruth"),
    ("sammie", "samantha", "samara"),
    ("sandi", "sandra", "cassandra", "sandy"),
    ("sibby", "sibyl", "sybil"),
    ("sissy", "cecelia", "cecilia"),
    ("stella", "estelle", "estrella"),
    ("sukey", "susan", "susannah"),
    ("suki", "susan", "suzette", "susannah"),
    ("sylvie", "sylvia", "silvana"),
    ("tabi", "tabitha", "tabby"),
    ("tallie", "natalie", "talia"),
    ("tessie", "theresa", "teresa", "tessa"),
    ("thea", "theodora", "dorothea", "theresa"),
    ("tibby", "isabella", "tibbie"),
    ("tillie", "matilda", "ottilie"),
    ("toby", "october", "tobitha"),
    ("tommie", "tamara", "tamsin"),
    ("toni", "antonia", "antoinette"),
    ("totie", "dorothy", "victoria"),
    ("tricia", "patricia", "beatrice"),
    ("trina", "katrina", "marina", "carolina"),
    ("trudie", "gertrude", "trudy"),
    ("una", "ursula", "unity"),
    ("vera", "veronica", "lavera"),
    ("vinnie", "lavinia", "virginia"),
    ("vita", "victoria", "davita"),
    ("vonnie", "yvonne", "lavonne"),
    ("wanda", "wenda", "gwendolyn"),
    ("xena", "alexena", "zena"),
    ("yola", "yolanda"),
    ("zelda", "griselda"),
    ("zena", "zenobia", "xena"),
    ("zia", "rosalia", "zenobia"),
    ("zosia", "sophia", "zoe"),
    # Additional male names
    ("abe", "abraham", "abner"),
    ("ace", "ace"),
    ("alf", "alfred", "alphonso"),
    ("archie", "archibald", "archer"),
    ("arnie", "arnold"),
    ("aug", "augustus", "augustine"),
    ("augie", "augustus", "augustine"),
    ("barney", "barnabas", "barnaby", "bernard"),
    ("basil", "basil"),
    ("bernie", "bernard", "bernardo"),
    ("bertie", "albert", "robert", "bert", "bertram"),
    ("biff", "william", "clifford"),
    ("birch", "birch"),
    ("bix", "bixby"),
    ("blaine", "blane"),
    ("blake", "blackwell"),
    ("bo", "robert", "beauregard"),
    ("bram", "abraham", "bramwell"),
    ("brice", "bryce"),
    ("bronson", "bronco"),
    ("brook", "brooks"),
    ("bucky", "buckley", "william"),
    ("butch", "william", "james"),
    ("cam", "cameron", "campbell"),
    ("carey", "carlisle", "charles"),
    ("cash", "cassius", "casper"),
    ("caz", "casimir", "caspar"),
    ("chad", "chadwick"),
    ("chipper", "charles"),
    ("cj", "christopher", "charles"),
    ("clem", "clement", "clementine"),
    ("cliff", "clifford", "clifton"),
    ("coby", "jacob", "coby"),
    ("con", "cornelius", "constantine"),
    ("cord", "cordell", "gordon"),
    ("cory", "cornelius", "corey"),
    ("cosmo", "cosimo"),
    ("curt", "curtis", "courtney"),
    ("dab", "dabney"),
    ("dawson", "david"),
    ("dex", "dexter"),
    ("dirk", "derek", "richard"),
    ("dolph", "rudolph", "adolph"),
    ("duke", "marmaduke", "stanford"),
    ("dunc", "duncan"),
    ("dwight", "dewight"),
    ("dyl", "dylan"),
    ("eli", "elijah", "elias", "elliot"),
    ("elliot", "elliot", "elias"),
    ("elmo", "elmwood"),
    ("emmet", "emmett", "emory"),
    ("erv", "ervin", "irving"),
    ("ev", "evan", "evander"),
    ("finn", "finnegan", "phineas"),
    ("flip", "philip", "phillip"),
    ("ford", "fordham", "stanford"),
    ("forrest", "forest"),
    ("fox", "foxworth"),
    ("gib", "gilbert", "gibson"),
    ("gordie", "gordon"),
    ("grace", "gracie", "gracey", "gray", "grey"),
    ("gray", "grayson"),
    ("ham", "hamilton", "hamlet"),
    ("harley", "harold", "harlan"),
    ("howie", "howard"),
    ("huck", "huckleberry"),
    ("iggy", "ignatius", "ignacio"),
    ("ira", "irving", "israel"),
    ("irv", "irving", "irvin"),
    ("izzy", "isidore", "israel", "isaiah"),
    ("jace", "jason", "jacoby"),
    ("jasper", "caspar"),
    ("jb", "james", "john"),
    ("jeb", "jebediah", "james"),
    ("jed", "jedidiah", "edgar"),
    ("jenks", "jenkins"),
    ("jesse", "jessamyn"),
    ("jett", "jethro"),
    ("jimbo", "james"),
    ("joab", "joab"),
    ("jock", "john", "james"),
    ("joel", "joel"),
    ("jonah", "jonas"),
    ("jordy", "jordan", "george"),
    ("jud", "judah", "judson"),
    ("kaz", "casimir", "kazimir"),
    ("kenji", "kenneth"),
    ("kev", "kevin"),
    ("lars", "lawrence", "laurence"),
    ("laz", "lazarus"),
    ("len", "leonard", "lennox"),
    ("lenny", "leonard", "lennox"),
    ("leo", "leonard", "leopold", "leon"),
    ("levi", "levi"),
    ("linc", "lincoln"),
    ("link", "lincoln"),
    ("luca", "lucas", "luke", "lucian"),
    ("luke", "lucas", "lucian"),
    ("mace", "mason", "macy"),
    ("manny", "manuel", "emmanuel", "immanuel"),
    ("marco", "mark", "marcus"),
    ("mars", "marcus", "marshall"),
    ("mort", "mortimer", "morton"),
    ("moss", "moses", "morris"),
    ("noel", "noel"),
    ("obie", "obadiah"),
    ("ogden", "ogden"),
    ("olin", "oliver", "olaf"),
    ("oz", "oswald", "oscar", "ozzy"),
    ("paddy", "patrick", "padraig"),
    ("pax", "paxton"),
    ("pip", "philip", "phillip"),
    ("raj", "rajesh", "rajan"),
    ("rand", "randall", "randolph", "randy"),
    ("reef", "reginald"),
    ("reid", "reid"),
    ("ren", "reginald", "renault"),
    ("rhys", "reece", "reese"),
    ("rowan", "rowland"),
    ("rowdy", "rowland"),
    ("royce", "royston"),
    ("rube", "rueben", "reuben"),
    ("rupert", "robert", "ruprecht"),
    ("rush", "rushford"),
    ("saul", "solomon"),
    ("scotty", "scott", "scottie"),
    ("seb", "sebastian"),
    ("sergei", "sergius"),
    ("shep", "shepherd", "stephen"),
    ("sig", "sigmund", "siegfried"),
    ("silas", "silvester"),
    ("soren", "sorin"),
    ("spence", "spencer"),
    ("tate", "tatham"),
    ("tex", "texas", "theodore"),
    ("thad", "thaddeus"),
    ("tobias", "toby"),
    ("tuck", "tucker"),
    ("ty", "tyler", "tyrone", "tyson"),
    ("ulric", "ulrich"),
    ("van", "vance", "vanderbilt"),
    ("vic", "victor", "vincent"),
    ("wade", "walden", "waldron"),
    ("ward", "edward", "howard", "wardell"),
    ("wash", "washington"),
    ("webb", "webber"),
    ("wiley", "william", "riley"),
    ("winton", "winston"),
    ("wolf", "wolfgang", "wolfe"),
    ("woody", "woodrow", "elwood"),
    ("wyatt", "wyat"),
    ("york", "yorkshire"),
    ("zane", "zachariah", "john"),
]

# Build flat lookup map — any name in a group finds all others in that group
NICKNAME_MAP = {}
for _group in _NICKNAME_GROUPS:
    _variants = list(dict.fromkeys(v.lower() for v in _group))
    for _name in _variants:
        if _name not in NICKNAME_MAP:
            NICKNAME_MAP[_name] = _variants
        else:
            NICKNAME_MAP[_name] = list(dict.fromkeys(NICKNAME_MAP[_name] + _variants))


POSITIONAL_NICKNAMES = {
    "junior": "jr",
    "deuce": "ii",
    "trey": "iii",
    "tres": "iii",
    "trip": "iii",
    "tripp": "iii",
    "triple": "iii",
    "quad": "iv",
    "quint": "v",
    "five": "v",
    "six": "vi",
}

NAME_SUFFIXES = {
    "jr", "jr.", "junior",
    "sr", "sr.", "senior",
    "ii", "iii", "iv", "v", "vi", "vii",
    "2nd", "3rd", "4th", "5th",
    "the second", "the third", "the fourth",
    "esq", "esq.", "esquire",
    "phd", "md", "dds", "dvm",
}


def strip_suffix(name: str) -> tuple[str, str | None]:
    """Returns (clean_name, suffix_found_or_None)."""
    parts = name.strip().split()
    if len(parts) > 1 and parts[-1].lower().rstrip(".") in NAME_SUFFIXES:
        return " ".join(parts[:-1]), parts[-1]
    if len(parts) > 2 and " ".join(parts[-2:]).lower() in NAME_SUFFIXES:
        return " ".join(parts[:-2]), " ".join(parts[-2:])
    return name.strip(), None


def get_name_variants(name: str) -> list:
    n = name.lower().strip()
    if n in NICKNAME_MAP:
        return NICKNAME_MAP[n]
    for variants in NICKNAME_MAP.values():
        if n in variants:
            return variants
    return [name]


def get_nb_token(nation_slug: str) -> str:
    secret_key = get_db().secrets.get_secret(scope="api", key="surus_server_nb_secret").value
    resp = requests.get(
        f"https://server.surusenterprises.com/auth/api_token/{nation_slug}",
        headers={"x-api-key": secret_key},
    )
    resp.raise_for_status()
    return resp.json()["access_token"]


def load_all_nations():
    if not WAREHOUSE_ID:
        return []
    try:
        result = get_db().statement_execution.execute_statement(
            warehouse_id=WAREHOUSE_ID,
            statement="SELECT `group`, slug, state FROM universal.nb.source_nation_table ORDER BY `group`",
            wait_timeout="30s",
        )
        if result.status.state != StatementState.SUCCEEDED:
            return []
        cols = [c.name for c in result.manifest.schema.columns]
        rows = (result.result.data_array if result.result else None) or []
        print(f"Loaded {len(rows)} nations")
        return [dict(zip(cols, row)) for row in rows]
    except Exception as e:
        print(f"Failed to load nations: {e}")
        return []

ALL_NATIONS = load_all_nations()


def ensure_log_table():
    if not WAREHOUSE_ID:
        return
    try:
        get_db().statement_execution.execute_statement(
            warehouse_id=WAREHOUSE_ID,
            statement="""
                CREATE TABLE IF NOT EXISTS universal.logging.contact_app_logs (
                    event_time  TIMESTAMP,
                    user_email  STRING,
                    user_name   STRING,
                    action      STRING,
                    nation_slug STRING,
                    details     STRING,
                    success     BOOLEAN,
                    error_message STRING
                )
            """,
            wait_timeout="30s",
        )
        print("Log table ready")
    except Exception as e:
        print(f"Could not ensure log table (non-critical): {e}")


def log_action(action: str, user_email: str, user_name: str,
               nation_slug: str = "", details: dict = None,
               success: bool = True, error_message: str = ""):
    def _write():
        if not WAREHOUSE_ID:
            return
        try:
            get_db().statement_execution.execute_statement(
                warehouse_id=WAREHOUSE_ID,
                statement="""
                    INSERT INTO universal.logging.contact_app_logs
                        (event_time, user_email, user_name, action, nation_slug, details, success, error_message)
                    VALUES
                        (current_timestamp(), :email, :uname, :action, :nation, :details,
                         CAST(:success AS BOOLEAN), :errmsg)
                """,
                parameters=[
                    StatementParameterListItem(name="email",  value=user_email or ""),
                    StatementParameterListItem(name="uname",  value=user_name or ""),
                    StatementParameterListItem(name="action", value=action),
                    StatementParameterListItem(name="nation", value=nation_slug or ""),
                    StatementParameterListItem(name="details", value=json.dumps(details or {})),
                    StatementParameterListItem(name="success", value="true" if success else "false"),
                    StatementParameterListItem(name="errmsg", value=error_message or ""),
                ],
                wait_timeout="15s",
            )
        except Exception as e:
            print(f"Log write failed (non-critical): {e}")
    threading.Thread(target=_write, daemon=True).start()


ensure_log_table()


@app.route("/search-nation")
@login_required
def search_nation():
    term = request.args.get("term", "").strip().lower()
    if not term:
        return jsonify({"success": False, "error": "Search term required"}), 400
    matches = [
        n for n in ALL_NATIONS
        if term in (n.get("group") or "").lower() or term in (n.get("slug") or "").lower()
    ]
    return jsonify({"success": True, "records": matches[:20]})


@app.route("/get-author-id")
@login_required
def get_author_id():
    nation_slug = request.args.get("nation_slug", "").strip()
    if not nation_slug or not WAREHOUSE_ID:
        return jsonify({"success": False, "nb_id": None})
    try:
        result = get_db().statement_execution.execute_statement(
            warehouse_id=WAREHOUSE_ID,
            statement="""
                SELECT nb_id, full_name, first_name, last_name
                FROM universal.prod.signups
                WHERE email = :email AND nation = :nation
                LIMIT 1
            """,
            parameters=[
                StatementParameterListItem(name="email", value=current_user.email),
                StatementParameterListItem(name="nation", value=nation_slug),
            ],
            wait_timeout="30s",
        )
        if result.status.state != StatementState.SUCCEEDED:
            return jsonify({"success": False, "nb_id": None})
        cols = [c.name for c in result.manifest.schema.columns]
        rows = (result.result.data_array if result.result else None) or []
        if rows:
            rec = dict(zip(cols, rows[0]))
            name = rec.get("full_name") or f"{rec.get('first_name','') or ''} {rec.get('last_name','') or ''}".strip()
            return jsonify({"success": True, "nb_id": rec.get("nb_id"), "name": name})
        return jsonify({"success": True, "nb_id": None})
    except Exception as e:
        return jsonify({"success": False, "error": str(e), "nb_id": None})


@app.route("/setup", methods=["GET", "POST"])
@login_required
def setup():
    if request.method == "POST":
        data = request.get_json() or {}
        nation_slug = data.get("nation_slug", "").strip()
        nation_name = data.get("nation_name", "").strip()
        author_nb_id = data.get("author_nb_id", "").strip()
        session["default_nation_slug"] = nation_slug
        session["default_nation_name"] = nation_name
        session["author_nb_id"] = author_nb_id
        log_action("nation_setup", current_user.email, current_user.name, nation_slug,
                   {"nation_name": nation_name, "author_nb_id": author_nb_id})
        return jsonify({"success": True})
    return render_template("setup.html",
                           default_nation_slug=session.get("default_nation_slug", ""),
                           default_nation_name=session.get("default_nation_name", ""),
                           author_nb_id=session.get("author_nb_id", ""))


@app.route("/search-by-name")
@login_required
def search_by_name():
    first = request.args.get("first", "").strip()
    last = request.args.get("last", "").strip()
    nation_slug = request.args.get("nation_slug", "").strip()

    if not last or not nation_slug:
        return jsonify({"success": False, "error": "Last name and nation required"}), 400
    if not WAREHOUSE_ID:
        return jsonify({"success": False, "error": "Warehouse not configured"}), 500

    first_clean, _ = strip_suffix(first)
    last_clean, _ = strip_suffix(last)
    first_lower = first_clean.lower()

    def run_query(statement, params):
        result = get_db().statement_execution.execute_statement(
            warehouse_id=WAREHOUSE_ID,
            statement=statement,
            parameters=params,
            wait_timeout="30s",
        )
        if result.status.state != StatementState.SUCCEEDED:
            raise Exception(result.status.error.message if result.status.error else "Query failed")
        cols = [c.name for c in result.manifest.schema.columns]
        rows = (result.result.data_array if result.result else None) or []
        return [dict(zip(cols, row)) for row in rows]

    select = """
        SELECT nb_id, first_name, last_name, full_name, suffix,
               COALESCE(`mailing_address.address1`, `registered_address.address1`, `home_address.address1`) AS address1,
               COALESCE(`mailing_address.city`, `registered_address.city`, `home_address.city`) AS city,
               COALESCE(`mailing_address.state`, `registered_address.state`, `home_address.state`) AS state,
               COALESCE(`mailing_address.zip`, `registered_address.zip`, `home_address.zip`) AS zip
        FROM universal.prod.signups
    """

    try:
        # Case 1: positional nickname (Trip, Trey, Deuce, Junior…) → search by last name + suffix
        if first_lower in POSITIONAL_NICKNAMES:
            suffix_val = POSITIONAL_NICKNAMES[first_lower]
            records = run_query(
                select + "WHERE LOWER(last_name) = LOWER(:last_name) AND LOWER(suffix) LIKE :suffix AND nation = :nation LIMIT 20",
                [
                    StatementParameterListItem(name="last_name", value=last_clean),
                    StatementParameterListItem(name="suffix", value=f"%{suffix_val}%"),
                    StatementParameterListItem(name="nation", value=nation_slug),
                ]
            )
            if records:
                return jsonify({"success": True, "records": records})

        # Case 2: normal nickname/name search with variants
        variants = get_name_variants(first_clean) if first_clean else [first_clean]
        conditions = " OR ".join([f"LOWER(first_name) = LOWER(:v{i})" for i in range(len(variants))])
        params = [StatementParameterListItem(name=f"v{i}", value=v) for i, v in enumerate(variants)]
        params += [
            StatementParameterListItem(name="last_name", value=last_clean),
            StatementParameterListItem(name="nation", value=nation_slug),
        ]
        records = run_query(
            select + f"WHERE ({conditions}) AND LOWER(last_name) = LOWER(:last_name) AND nation = :nation LIMIT 20",
            params
        )
        if records:
            return jsonify({"success": True, "records": records})

        # Case 3: fallback — last name only (catches unusual nicknames not in our map)
        records = run_query(
            select + "WHERE LOWER(last_name) = LOWER(:last_name) AND nation = :nation LIMIT 20",
            [
                StatementParameterListItem(name="last_name", value=last_clean),
                StatementParameterListItem(name="nation", value=nation_slug),
            ]
        )
        return jsonify({"success": True, "records": records, "fallback": True})

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/search-signup")
@login_required
def search_signup():
    name = request.args.get("name", "").strip()
    nation_slug = request.args.get("nation_slug", "").strip()

    if not name or not nation_slug:
        return jsonify({"success": False, "error": "Name and nation slug are required"}), 400

    if not WAREHOUSE_ID:
        return jsonify({"success": False, "error": "DATABRICKS_WAREHOUSE_ID not configured"}), 500

    try:
        result = get_db().statement_execution.execute_statement(
            warehouse_id=WAREHOUSE_ID,
            statement="""
                SELECT nb_id, first_name, last_name, full_name,
                       COALESCE(`mailing_address.address1`, `registered_address.address1`, `home_address.address1`) AS address1,
                       COALESCE(`mailing_address.city`, `registered_address.city`, `home_address.city`) AS city,
                       COALESCE(`mailing_address.state`, `registered_address.state`, `home_address.state`) AS state,
                       COALESCE(`mailing_address.zip`, `registered_address.zip`, `home_address.zip`) AS zip
                FROM universal.prod.signups
                WHERE full_name ILIKE :term
                  AND nation = :nation
                LIMIT 20
            """,
            parameters=[
                StatementParameterListItem(name="term", value=f"{strip_suffix(name)[0]}%"),
                StatementParameterListItem(name="nation", value=nation_slug),
            ],
            wait_timeout="30s",
        )

        if result.status.state != StatementState.SUCCEEDED:
            msg = result.status.error.message if result.status.error else "Query failed"
            return jsonify({"success": False, "error": msg}), 500

        cols = [c.name for c in result.manifest.schema.columns]
        rows = result.result.data_array or []
        records = [dict(zip(cols, row)) for row in rows]

        return jsonify({"success": True, "records": records})

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")


def parse_image_with_ai(raw: bytes, filename: str) -> pd.DataFrame:
    mime_map = {
        "jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png",
        "gif": "image/gif", "bmp": "image/bmp", "webp": "image/webp", "tiff": "image/tiff",
    }
    ext = filename.lower().rsplit(".", 1)[-1]
    mime = mime_map.get(ext, "image/png")
    b64 = base64.b64encode(raw).decode("utf-8")
    resp = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={"Authorization": f"Bearer {OPENROUTER_API_KEY}", "Content-Type": "application/json"},
        json={
            "model": "openai/gpt-4o-mini",
            "messages": [{
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                    {"type": "text", "text": (
                        "Extract all contact or person data visible in this image into a JSON array of objects. "
                        "Each person/row should be one object with column headers as keys. "
                        "Include every column you can see: name, phone, email, address, date contacted, notes, IDs, etc. "
                        "Return ONLY a valid JSON array — no markdown, no code fences, no explanation."
                    )},
                ],
            }],
        },
        timeout=90,
    )
    resp.raise_for_status()
    content = resp.json()["choices"][0]["message"]["content"].strip()
    if content.startswith("```"):
        content = content.split("```")[1]
        if content.startswith("json"):
            content = content[4:]
    data = json.loads(content.strip())
    if isinstance(data, list) and data:
        return pd.DataFrame(data)
    return pd.DataFrame({"extracted": ["No structured data found in image"]})


def parse_upload(file) -> pd.DataFrame:
    name = file.filename.lower()
    raw = file.read()
    if name.endswith(".csv") or name.endswith(".txt"):
        for kwargs in [
            {"sep": ",", "quotechar": '"', "engine": "python"},
            {"sep": ",", "quotechar": '"', "quoting": 0, "on_bad_lines": "skip"},
            {"sep": "\t", "engine": "python"},
            {"sep": "|", "engine": "python"},
            {"sep": ";", "engine": "python"},
            {"sep": ",", "on_bad_lines": "skip"},
        ]:
            try:
                df = pd.read_csv(io.BytesIO(raw), **kwargs)
                if len(df.columns) > 1:
                    return df
            except Exception:
                continue
        return pd.read_csv(io.BytesIO(raw), on_bad_lines="skip")
    elif name.endswith(".xlsx") or name.endswith(".xls"):
        return pd.read_excel(io.BytesIO(raw))
    elif name.endswith(".docx"):
        from docx import Document
        doc = Document(io.BytesIO(raw))
        for table in doc.tables:
            headers = [c.text.strip() for c in table.rows[0].cells]
            rows = [[c.text.strip() for c in r.cells] for r in table.rows[1:]]
            if headers:
                return pd.DataFrame(rows, columns=headers)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return pd.DataFrame({"text": text.splitlines()})
    elif name.endswith(".json"):
        data = json.loads(raw.decode("utf-8"))
        if isinstance(data, list):
            return pd.DataFrame(data)
        if isinstance(data, dict):
            for v in data.values():
                if isinstance(v, list):
                    return pd.DataFrame(v)
        return pd.DataFrame([data])
    elif name.endswith(".pdf"):
        try:
            import pdfplumber
            rows = []
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                for page in pdf.pages:
                    for table in (page.extract_tables() or []):
                        if table:
                            headers = table[0]
                            for row in table[1:]:
                                rows.append(dict(zip(headers, row)))
            if rows:
                return pd.DataFrame(rows)
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                text = "\n".join(p.extract_text() or "" for p in pdf.pages)
            return pd.DataFrame({"text": [l for l in text.splitlines() if l.strip()]})
        except ImportError:
            return pd.DataFrame({"error": ["Install pdfplumber to parse PDFs: pip install pdfplumber"]})
    elif any(name.endswith(f".{ext}") for ext in ("png", "jpg", "jpeg", "gif", "bmp", "webp", "tiff")):
        return parse_image_with_ai(raw, file.filename)
    else:
        try:
            return pd.read_csv(io.BytesIO(raw))
        except Exception:
            text = raw.decode("utf-8", errors="ignore")
            return pd.DataFrame({"text": text.splitlines()})


def _apply_mapping_locally(column_mapping: dict, all_rows: list) -> list:
    methods_map = {m.replace("_", " "): m for m in CONTACT_METHODS}
    methods_map.update({m: m for m in CONTACT_METHODS})
    statuses_map = {s.replace("_", " "): s for s in CONTACT_STATUSES}
    statuses_map.update({s: s for s in CONTACT_STATUSES})
    date_fmts = ("%Y-%m-%d", "%m/%d/%Y", "%m/%d/%y", "%d/%m/%Y",
                 "%B %d, %Y", "%b %d, %Y", "%Y/%m/%d")
    result = []
    for raw in all_rows:
        row = {}
        for src, nb in column_mapping.items():
            if not nb:
                continue
            val = str(raw.get(src, "") or "").strip()
            if not val:
                continue
            if nb in ("signup_id", "author_id"):
                digits = "".join(c for c in val if c.isdigit())
                if digits:
                    row[nb] = digits
            elif nb == "contact_method":
                key = val.lower().replace("-", "_").replace(" ", "_")
                row[nb] = methods_map.get(val.lower(), methods_map.get(key, key))
            elif nb == "contact_status":
                key = val.lower().replace("-", "_").replace(" ", "_")
                row[nb] = statuses_map.get(val.lower(), statuses_map.get(key, key))
            elif nb == "contact_date":
                parsed = None
                for fmt in date_fmts:
                    try:
                        parsed = datetime.strptime(val, fmt).strftime("%Y-%m-%d")
                        break
                    except Exception:
                        pass
                row[nb] = parsed or val
            elif nb == "_full_name":
                parts = val.split(None, 1)
                row["_first_name"] = parts[0] if parts else ""
                row["_last_name"] = parts[1] if len(parts) > 1 else ""
            else:
                row[nb] = val
        result.append(row)
    return result


def ai_map_and_clean(columns: list, all_rows: list) -> dict:
    sample = all_rows[:6]
    prompt = f"""Map these source columns to NationBuilder contact fields.

Valid NB fields:
- signup_id: NationBuilder person ID (digits only)
- author_id: NationBuilder ID of who logged the contact (digits only)
- contact_method: one of {json.dumps(CONTACT_METHODS)}
- contact_status: one of {json.dumps(CONTACT_STATUSES)}
- contact_date: date contacted, normalize to YYYY-MM-DD
- content: free-text notes

Name handling (IMPORTANT — do NOT mark name columns as null):
- Full name column (e.g. "Name", "Full Name", "Contact") → use "_full_name"
- First name only → use "_first_name"
- Last name only → use "_last_name"

Source columns: {json.dumps(columns)}
Sample rows: {json.dumps(sample)}

Return ONLY JSON (no markdown):
{{"column_mapping": {{"source_column": "nb_field_or_null"}}, "notes": "one line summary"}}"""

    resp = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={"Authorization": f"Bearer {OPENROUTER_API_KEY}", "Content-Type": "application/json"},
        json={"model": "openai/gpt-4o-mini", "messages": [{"role": "user", "content": prompt}]},
        timeout=30,
    )
    resp.raise_for_status()
    content = resp.json()["choices"][0]["message"]["content"].strip()
    if content.startswith("```"):
        content = content.split("```")[1]
        if content.startswith("json"):
            content = content[4:]
    mapping = json.loads(content.strip())
    mapping["cleaned_rows"] = _apply_mapping_locally(mapping.get("column_mapping", {}), all_rows)
    return mapping


@app.route("/bulk")
@login_required
def bulk():
    return render_template("bulk.html",
                           default_nation_slug=session.get("default_nation_slug", ""),
                           default_nation_name=session.get("default_nation_name", ""),
                           author_nb_id=session.get("author_nb_id", ""))


@app.route("/bulk/upload", methods=["POST"])
@login_required
def bulk_upload():
    if "file" not in request.files or request.files["file"].filename == "":
        return jsonify({"success": False, "error": "No file uploaded"}), 400
    try:
        df = parse_upload(request.files["file"])
        df = df.dropna(how="all").head(500)
        columns = list(df.columns)
        all_rows_raw = df.fillna("").astype(str).to_dict(orient="records")
        result = ai_map_and_clean(columns, all_rows_raw)
        cleaned = result.get("cleaned_rows", [])
        return jsonify({
            "success": True,
            "columns": columns,
            "mapping": result,
            "preview": cleaned[:10],
            "total_rows": len(cleaned),
            "all_rows": cleaned,
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/bulk/import", methods=["POST"])
@login_required
def bulk_import():
    data = request.get_json()
    nation_slug = data.get("nation_slug", "").strip()
    rows = data.get("rows", [])
    imported_by = data.get("imported_by", "").strip()
    if not nation_slug:
        return jsonify({"success": False, "error": "Nation slug required"}), 400
    if not rows:
        return jsonify({"success": False, "error": "No rows to import"}), 400
    try:
        token = get_nb_token(nation_slug)
    except Exception as e:
        return jsonify({"success": False, "error": f"Auth failed: {e}"}), 500

    now = datetime.now(timezone.utc)
    timestamp = now.strftime("%B %d, %Y at %I:%M:%S %p UTC")
    import_tag = f"\n\n--- Bulk import by: {imported_by} | {timestamp} ---" if imported_by else f"\n\n--- Bulk import | {timestamp} ---"

    url = f"https://{nation_slug}.nationbuilder.com/api/v2/contacts"
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
        "Accept": "application/json",
    }
    results = {"success": 0, "failed": 0, "errors": []}
    for i, row in enumerate(rows):
        attributes = {k: v for k, v in row.items()
                      if k in ("contact_method", "contact_status", "content")}

        # Build content: date contacted (if present) → user notes → import stamp
        parts = []
        contact_date = str(row.get("contact_date", "") or "").strip()
        if contact_date:
            try:
                formatted_date = datetime.strptime(contact_date, "%Y-%m-%d").strftime("%B %d, %Y")
            except Exception:
                formatted_date = contact_date
            parts.append(f"Date Contacted: {formatted_date}")
        existing_content = attributes.get("content", "") or ""
        if existing_content:
            parts.append(existing_content)
        parts.append(import_tag.strip())
        attributes["content"] = "\n\n".join(parts)
        relationships = {}
        for field, rel in [("signup_id", "signup"), ("author_id", "author")]:
            if row.get(field):
                relationships[rel] = {"data": {"type": "signups", "id": str(row[field])}}
        body = {"data": {"type": "contacts", "attributes": attributes}}
        if relationships:
            body["data"]["relationships"] = relationships
        try:
            resp = requests.post(url, headers=headers, json=body)
            resp.raise_for_status()
            results["success"] += 1
        except requests.HTTPError as e:
            results["failed"] += 1
            detail = e.response.json() if e.response else str(e)
            results["errors"].append({"row": i + 1, "error": str(e), "detail": detail})
        except Exception as e:
            results["failed"] += 1
            results["errors"].append({"row": i + 1, "error": str(e)})
    log_action("bulk_import", current_user.email, current_user.name, nation_slug, {
        "total_rows": len(rows),
        "imported": results["success"],
        "failed": results["failed"],
        "imported_by": imported_by or current_user.name,
    }, success=results["failed"] == 0)
    return jsonify({"success": True, "results": results})


@app.route("/login")
def login():
    app_url = os.getenv("APP_URL", "http://localhost:5000")
    redirect_uri = app_url + "/auth/callback"
    return google.authorize_redirect(redirect_uri)

@app.route("/auth/callback")
def auth_callback():
    token = google.authorize_access_token()
    userinfo = token.get("userinfo") or google.userinfo()
    if not userinfo:
        return render_template("login.html", error="Could not retrieve account info from Google.")
    email = userinfo.get("email", "")
    if not email.endswith("@surusenterprises.com"):
        return render_template("login.html", error="Access restricted to Surus Enterprises accounts only.")
    user = User(
        id=email,
        email=email,
        name=userinfo.get("name", email),
        picture=userinfo.get("picture", ""),
    )
    _users[email] = user
    login_user(user, remember=True)
    log_action("login", email, userinfo.get("name", email))
    return redirect("/setup")

@app.route("/logout")
@login_required
def logout():
    logout_user()
    return redirect("/login")


@app.route("/")
@login_required
def index():
    return render_template("index.html",
                           contact_methods=CONTACT_METHODS,
                           contact_statuses=CONTACT_STATUSES,
                           default_nation_slug=session.get("default_nation_slug", ""),
                           default_nation_name=session.get("default_nation_name", ""),
                           author_nb_id=session.get("author_nb_id", ""))


@app.route("/import", methods=["POST"])
@login_required
def import_contact():
    form = request.form
    nation_slug = form.get("nation_slug", "").strip()

    if not nation_slug:
        return jsonify({"success": False, "error": "Nation slug is required"}), 400

    attributes = {}

    relationships = {}

    for field, rel_type in [("author_id", "author"), ("signup_id", "signup"),
                             ("broadcaster_id", "broadcaster"), ("path_id", "path")]:
        val = form.get(field, "").strip()
        if val:
            relationships[rel_type] = {"data": {"type": "signups", "id": val}}

    path_step = form.get("path_step_id", "").strip()
    if path_step:
        relationships["path_step"] = {"data": {"type": "path_steps", "id": path_step}}

    # Build content: date contacted (if provided) → user notes
    content = form.get("content", "").strip()
    contact_date = form.get("contact_date", "").strip()
    content_parts = []
    if contact_date:
        try:
            formatted_date = datetime.strptime(contact_date, "%Y-%m-%d").strftime("%B %d, %Y")
        except Exception:
            formatted_date = contact_date
        content_parts.append(f"Date Contacted: {formatted_date}")
    if content:
        content_parts.append(content)
    if content_parts:
        attributes["content"] = "\n\n".join(content_parts)

    for field in ["contact_status", "contact_method"]:
        val = form.get(field, "").strip()
        if val:
            attributes[field] = val

    pc = form.get("pc_in_cents", "").strip()
    if pc:
        try:
            attributes["pc_in_cents"] = int(pc)
        except ValueError:
            return jsonify({"success": False, "error": "Political capital must be a whole number"}), 400

    body = {"data": {"type": "contacts", "attributes": attributes}}
    if relationships:
        body["data"]["relationships"] = relationships

    try:
        token = get_nb_token(nation_slug)
        url = f"https://{nation_slug}.nationbuilder.com/api/v2/contacts"
        resp = requests.post(
            url,
            headers={
                "Authorization": f"Bearer {token}",
                "Content-Type": "application/json",
                "Accept": "application/json",
            },
            json=body,
        )
        resp.raise_for_status()
        data = resp.json()
        log_action("single_import", current_user.email, current_user.name, nation_slug, {
            "contact_id": (data.get("data") or {}).get("id"),
            "signup_id": form.get("signup_id", ""),
            "method": form.get("contact_method", ""),
            "status": form.get("contact_status", ""),
        })
        return jsonify({"success": True, "data": data})
    except requests.HTTPError as e:
        detail = None
        if e.response is not None:
            try:
                detail = e.response.json()
            except Exception:
                detail = e.response.text
        log_action("single_import", current_user.email, current_user.name, nation_slug,
                   {"signup_id": form.get("signup_id", "")},
                   success=False, error_message=str(e))
        return jsonify({"success": False, "error": str(e), "detail": detail}), 400
    except Exception as e:
        log_action("single_import", current_user.email, current_user.name, nation_slug,
                   {}, success=False, error_message=str(e))
        return jsonify({"success": False, "error": str(e)}), 500


if __name__ == "__main__":
    app.run(debug=True)
